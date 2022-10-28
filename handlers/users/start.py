from aiogram import types
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.builtin import CommandStart
from aiogram.types import CallbackQuery, InputFile, ReplyKeyboardMarkup, ReplyKeyboardRemove
from aiogram.dispatcher.filters import Regexp
import os

from data.config import ADMINS
from keyboards.default.check_menu import check_menu
from keyboards.inline.menu import menu, about_key, about_bridge, play
from loader import dp,bot

from docx import Document
from docx.shared import Inches
phone_regex = '^[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}$'

@dp.message_handler(CommandStart())
async def bot_start(message: types.Message):
    await message.answer(f"Ҳурматли фойдаланувчи сиз бу бот орқали бизнинг"
                         f" корхонамизга ишга кириш учун ариза қолдириб,ўз "
                         f"резюменгизни юборишингиз мумкин!",reply_markup=menu)


@dp.callback_query_handler(text = 'send' )
async def anketa(call:CallbackQuery,state:FSMContext):
    await call.message.delete()
    await call.message.answer('Фамилия, исм-шарифингиз')
    await state.set_state('ism')

@dp.message_handler(state='ism')
async def ism_def(message:types.Message,state:FSMContext):
    ism = message.text
    await message.answer('Тугилган санангиз, 03.05.1997')
    await state.update_data(
        {
            'ism':ism
        }
    )
    await state.set_state('birthdate')
@dp.message_handler(state='birthdate')
async def birthdate_def(message:types.Message,state:FSMContext):
    birthdate = message.text
    await state.update_data(
        {
            'birthdate':birthdate
        }
    )
    await message.answer('Яшаш манзилингиз (Шахар, Туман)')
    await state.set_state('place')

@dp.message_handler(state='place')
async def phone_def(message:types.Message,state:FSMContext):
    joy = message.text
    await state.update_data(
        {
            'joy':joy
        }
    )
    await message.answer('Телефон ракамингиз.+998901234567')
    await state.set_state('phone')

@dp.message_handler(Regexp(phone_regex),state = 'phone')
async def about_def(message:types.Message,state:FSMContext):
    phone = message.text
    await state.update_data(
        {
            'phone':phone
        }
    )
    await message.answer('Маълумотингиз?',reply_markup=about_key)
    await state.set_state('about')
@dp.message_handler(state = 'phone')
async def space_def(message:types.Message):
    await message.answer('🛑 Telefon raqamingizni kiriting!')
@dp.callback_query_handler(state='about')
async def unver_def(call:CallbackQuery,state:FSMContext):
    unver = call.data
    await state.update_data(
          {
            'unver':unver
          }
            )

    await call.message.answer('Оилавий аxволингиз?',reply_markup=about_bridge)
    await state.set_state('male_female')
    await call.message.delete()

@dp.callback_query_handler(state='male_female')
async def male_female_def(call:CallbackQuery,state:FSMContext):
    await call.message.delete()
    male_female = call.data
    await state.update_data(
           {
            'male_female':male_female
          }
          )

    await call.message.answer('Қайси компьютер дастурларида ишлагансиз?')
    await state.set_state('computer')

@dp.message_handler(state = 'computer')
async def sud_def(message:types.Message,state:FSMContext):
    computer = message.text
    await state.update_data(
        {
            'computer':computer
        }
    )
    await message.answer('Охирги маротаба каерларда ишлагансиз?')
    await state.set_state('worked')

@dp.message_handler(state  = 'worked')
async def interests_def(message:types.Message,state:FSMContext):
    worked = message.text
    await state.update_data(
        {
            'worked':worked
        }
    )

    await message.answer('Охирги иш урнингиздаги ойлик маошингиз канча булган')
    await state.set_state('money')

@dp.message_handler(state = 'money')
async def team_def(message:types.Message,state:FSMContext):
    money = message.text
    await state.update_data(
        {
            'money':money
        }
    )
    await message.answer('Бизда канча микдорли маошга ишламокчисиз (ёзишингиз шарт)')
    await state.set_state('we_money')
@dp.message_handler(state='we_money')
async def computer_def(message:types.Message,state:FSMContext):
    we_money = message.text
    await state.update_data(
        {
            'we_money':we_money
        }
    )

    await message.answer('Кайси йуналишда (оддий ходим, сотув, бошкарув, хисоб китоб, '
                         'техник, слесарь, сварщик,бухгалтерия) бемалол кийналмай ишлай'
                         ' оласиз?')
    await state.set_state('yunalish')

@dp.message_handler(state='yunalish')
async def yunalish_def(message:types.Message,state:FSMContext):
    xodim = message.text
    await state.update_data(
        {
            'xodim':xodim
        }
    )

    await message.answer('Тасдиқланг',reply_markup=play)
    await state.set_state('play')

@dp.callback_query_handler(state='play')
async def play_def(call:CallbackQuery,state:FSMContext):
    await call.message.delete()
    data = await state.get_data()
    ism = data.get('ism')
    birthdate=data.get('birthdate')
    joy = data.get('joy')
    phone = data.get('phone')
    unver = data.get('unver')
    male_female = data.get('male_female')
    computer = data.get('computer')
    worked = data.get('worked')
    money = data.get('money')
    we_money = data.get('we_money')
    xodim = data.get('xodim')

    document = Document()

    document.add_heading('Malumotlar', 0)
    p = document.add_paragraph()
    p.add_run('Ma\'lumotlar maxfiy saqlanadi!').bold = True

    document.add_paragraph(f'Ism: {ism} ',style='List Number')
    document.add_paragraph(f'Tug\'ilgan sana: {birthdate}',style='List Number')
    document.add_paragraph(f'Turar joy: {joy}',style='List Number')
    document.add_paragraph(f'Telefon raqami: {phone} ',style='List Number')
    document.add_paragraph(f'Ta\'lim: {unver}',style='List Number')
    document.add_paragraph(f'Turmush Tarzi: {male_female}',style='List Number')
    document.add_paragraph(f'Kompyuter Dasturlari: {computer} ',style='List Number')
    document.add_paragraph(f'Avvalgi ish joyi: {worked} ',style='List Number')
    document.add_paragraph(f'Qancha maoshga ishlagani: {money} ',style='List Number')
    document.add_paragraph(f'Bizda qancha maoshga ishlashi: {we_money} ',style='List Number')
    document.add_paragraph(f'Qanday xodim bo\'lib ishlashi: {xodim} ',style='List Number')
    document.add_page_break()
    document.save(f'{call.from_user.id}.docx')

    doc = InputFile(f'{call.from_user.id}.docx', filename=f'{call.from_user.full_name}.docx')
    await call.message.answer_document(document=doc,caption='Bu sizning resumingiz')
    await call.message.answer('Operatorga yuborish uchun pastdagi tugmani bosing!',reply_markup=check_menu)
    await state.finish()

@dp.message_handler(text = '✳️ Operatorga yuborish')
async def bbb(message:types.Message):
    doc = InputFile(f'{message.from_user.id}.docx', filename=f'{message.from_user.full_name}.docx')
    os.remove(path=f'{message.from_user.id}.docx')
    await bot.send_document(chat_id=ADMINS[0],document=doc,caption=f"@{message.from_user.username}")
    await message.answer('<b>Ma\'lumotlaringiz yuborildi,\n\n'
                         '/start buyrug\'i orqali botni ishga tushiring.</b>',reply_markup=ReplyKeyboardRemove())



